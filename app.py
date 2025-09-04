import io
import base64
import hashlib
from typing import Optional, Any

import streamlit as st

# =========================
# PAGE SETUP
# =========================
st.set_page_config(page_title="PDF → DOCX (Azure DI) + Credits", page_icon="📄", layout="wide")
st.title("📄 PDF → DOCX with Azure Document Intelligence (Read)")
st.caption("Upload a PDF → Azure DI (prebuilt-read) extracts text → Download a .docx • Pricing: ₹3 per page (3 credits)")

# =========================
# PRICING / CONSTANTS
# =========================
CREDITS_START_BALANCE = 10_000
PRICE_PER_PAGE_CREDITS = 3  # ₹3 == 3 credits

# =========================
# SECRETS / CONFIG
# =========================
def get_secret(key: str, default: Optional[str] = None) -> Optional[str]:
    try:
        return st.secrets[key]  # type: ignore[attr-defined]
    except Exception:
        return default

AZURE_DI_ENDPOINT = get_secret("AZURE_DI_ENDPOINT")
AZURE_DI_KEY = get_secret("AZURE_DI_KEY")
# Use st.secrets for admin PIN; default to "1133344444" if not provided
ADMIN_PIN = str(get_secret("ADMIN_PIN", "1133344444"))

# =========================
# STATE INIT
# =========================
# Wallet (locked by default; only admin can refill)
if "credits_balance" not in st.session_state:
    st.session_state.credits_balance = CREDITS_START_BALANCE

# Prevent double-charging the same file during re-runs
if "charged_docs" not in st.session_state:
    st.session_state.charged_docs = {}  # {file_hash: {"pages": int, "cost": int}}

# Track last transaction
if "last_txn" not in st.session_state:
    st.session_state.last_txn = None    # {"file": str, "pages": int, "cost": int, "ts": "..."}

# Track admin auth
if "is_admin" not in st.session_state:
    st.session_state.is_admin = False

# =========================
# AZURE SDK IMPORTS
# =========================
try:
    from azure.ai.documentintelligence import DocumentIntelligenceClient
    from azure.core.credentials import AzureKeyCredential
except Exception:
    DocumentIntelligenceClient = None
    AzureKeyCredential = None

try:
    from azure.ai.documentintelligence.models import AnalyzeDocumentRequest  # type: ignore
except Exception:
    AnalyzeDocumentRequest = None  # type: ignore

from docx import Document
from docx.shared import Pt
from datetime import datetime

# =========================
# SIDEBAR: CREDITS + ADMIN
# =========================
with st.sidebar:
    st.subheader("💳 Credits")
    max_display = max(CREDITS_START_BALANCE, st.session_state.credits_balance)
    pct = min(max(st.session_state.credits_balance / float(max_display), 0.0), 1.0)
    st.progress(pct, text=f"Balance: {st.session_state.credits_balance} credits")

    # Last transaction (pretty card)
    if st.session_state.last_txn:
        txn = st.session_state.last_txn
        st.markdown(
            f"""
            <div style="
                background:#f5f8ff;
                padding:12px;
                border-radius:10px;
                border:1px solid #d1e3ff;
                margin-top:12px;
            ">
              <div style="font-weight:600;color:#1f4396;margin-bottom:6px;">🧾 Last Transaction</div>
              <div style="font-size:13px;line-height:1.4;">
                <div><b>File:</b> {txn['file']}</div>
                <div><b>Pages:</b> {txn['pages']}</div>
                <div><b>Credits:</b> {txn['cost']} (₹{txn['cost']})</div>
                <div style="color:#666;"><b>Time:</b> {txn.get('ts','')}</div>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    # Admin panel (refill only here)
    with st.expander("🔐 Admin Panel", expanded=False):
        pin_entered = st.text_input("Enter Admin PIN", type="password")
        if st.button("Login"):
            # Compare strictly to st.secrets ADMIN_PIN
            if pin_entered == ADMIN_PIN:
                st.session_state.is_admin = True
                st.success("Login successful.")
            else:
                st.session_state.is_admin = False
                st.error("Wrong password")

        if st.session_state.is_admin:
            st.markdown("**Admin Controls**")
            topup_amt = st.number_input("Top-up amount (credits)", min_value=1, value=100, step=50)
            if st.button("Top-up Wallet"):
                st.session_state.credits_balance += int(topup_amt)
                st.success(f"Wallet topped up by {int(topup_amt)} credits.")

            # Optional: Admin reset (credits + history)
            if st.button("Reset Wallet & History"):
                st.session_state.credits_balance = CREDITS_START_BALANCE
                st.session_state.charged_docs.clear()
                st.session_state.last_txn = None
                st.success("Wallet reset to 10,000 and history cleared.")

# =========================
# SETTINGS / KEYS
# =========================
with st.expander("⚙️ Settings", expanded=False):
    add_page_breaks = st.checkbox("Insert page breaks between PDF pages", value=True)
    include_confidence = st.checkbox("Append line confidence (debug)", value=False)

# If secrets not configured, allow user to input for this run
if not AZURE_DI_ENDPOINT or not AZURE_DI_KEY:
    st.info("Azure DI endpoint/key not found in st.secrets. Enter them for this session.")
    AZURE_DI_ENDPOINT = st.text_input("AZURE_DI_ENDPOINT", AZURE_DI_ENDPOINT or "", placeholder="https://<resourcename>.cognitiveservices.azure.com/")
    AZURE_DI_KEY = st.text_input("AZURE_DI_KEY", AZURE_DI_KEY or "", type="password")

uploaded = st.file_uploader("Upload a PDF", type=["pdf"], accept_multiple_files=False)

# =========================
# HELPERS
# =========================
@st.cache_resource(show_spinner=False)
def make_client(endpoint: str, key: str):
    if DocumentIntelligenceClient is None or AzureKeyCredential is None:
        raise RuntimeError("Azure SDK not installed. Run: pip install azure-ai-documentintelligence python-docx")
    if not endpoint or not key:
        raise RuntimeError("Missing Azure DI endpoint or key.")
    return DocumentIntelligenceClient(endpoint=endpoint, credential=AzureKeyCredential(key))

def analyze_pdf_bytes(client: Any, pdf_bytes: bytes):
    """
    Calls Azure Document Intelligence 'prebuilt-read' across multiple SDK variants.
    Tries (in order):
      A) document=pdf_bytes, content_type="application/pdf"         (newer SDKs)
      B) body={"base64Source": "..."}                               (older SDKs)
      C) body=AnalyzeDocumentRequest(bytes_source=pdf_bytes)        (some mid SDKs)
    """
    last_err = None

    # A) Newer SDK
    try:
        poller = client.begin_analyze_document(
            model_id="prebuilt-read",
            document=pdf_bytes,
            content_type="application/pdf",
        )
        return poller.result()
    except Exception as e:
        last_err = e

    # B) Older SDK (base64 body)
    try:
        b64 = base64.b64encode(pdf_bytes).decode("utf-8")
        poller = client.begin_analyze_document(
            "prebuilt-read",
            body={"base64Source": b64},
        )
        return poller.result()
    except Exception as e:
        last_err = e

    # C) Typed request model
    try:
        if AnalyzeDocumentRequest is not None:
            req = AnalyzeDocumentRequest(bytes_source=pdf_bytes)  # type: ignore
            poller = client.begin_analyze_document(
                model_id="prebuilt-read",
                body=req,
                content_type="application/pdf",
            )
            return poller.result()
    except Exception as e:
        last_err = e

    raise last_err

def result_to_docx_bytes(result, insert_page_breaks: bool = True, show_conf: bool = False) -> bytes:
    """Convert DI 'prebuilt-read' result into a .docx."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if not getattr(result, "pages", None):
        doc.add_paragraph(getattr(result, "content", "") or "No content found.")
    else:
        for idx, page in enumerate(result.pages):
            doc.add_heading(f"Page {idx+1}", level=2)
            if getattr(page, "lines", None):
                for ln in page.lines:
                    text = ln.content or ""
                    if show_conf and hasattr(ln, "spans") and ln.spans:
                        try:
                            confs = [getattr(s, "confidence", None) for s in ln.spans if getattr(s, "confidence", None) is not None]
                            if confs:
                                text += f"  [conf~{sum(confs)/len(confs):.2f}]"
                        except Exception:
                            pass
                    if text.strip():
                        doc.add_paragraph(text)
            else:
                # Fallback: paragraphs
                paras = []
                for p in getattr(result, "paragraphs", []) or []:
                    if getattr(p, "spans", None):
                        if any(getattr(sp, "offset", None) is not None for sp in p.spans):
                            paras.append(p.content)
                if paras:
                    for p in paras:
                        doc.add_paragraph(p)
                else:
                    doc.add_paragraph(getattr(result, "content", "") or "")

            if insert_page_breaks and idx < len(result.pages) - 1:
                doc.add_page_break()

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def file_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()

def charge_credits_once(file_id: str, pages: int, filename: str) -> int:
    """
    Deduct credits exactly once per unique file hash.
    Returns the charged cost (0 if already charged).
    """
    cost = pages * PRICE_PER_PAGE_CREDITS

    # Already charged?
    if file_id in st.session_state.charged_docs:
        st.info("⚠️ This file was already processed earlier. No credits deducted again.")
        return 0

    # Sufficient balance?
    if st.session_state.credits_balance < cost:
        raise RuntimeError(
            f"Insufficient credits: need {cost}, have {st.session_state.credits_balance}. "
            "Please ask an admin to top-up credits."
        )

    # Deduct and record
    st.session_state.credits_balance -= cost
    st.session_state.charged_docs[file_id] = {"pages": pages, "cost": cost}
    st.session_state.last_txn = {
        "file": filename,
        "pages": pages,
        "cost": cost,
        "ts": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    return cost

# =========================
# MAIN FLOW
# =========================
uploaded = st.file_uploader("Upload a PDF", type=["pdf"], accept_multiple_files=False)

# If secrets not configured, allow user to input for this run
with st.expander("⚙️ Settings", expanded=False):
    add_page_breaks = st.checkbox("Insert page breaks between PDF pages", value=True)
    include_confidence = st.checkbox("Append line confidence (debug)", value=False)

if not AZURE_DI_ENDPOINT or not AZURE_DI_KEY:
    st.info("Azure DI endpoint/key not found in st.secrets. Enter them for this session.")
    AZURE_DI_ENDPOINT = st.text_input("AZURE_DI_ENDPOINT", AZURE_DI_ENDPOINT or "", placeholder="https://<resourcename>.cognitiveservices.azure.com/")
    AZURE_DI_KEY = st.text_input("AZURE_DI_KEY", AZURE_DI_KEY or "", type="password")

if uploaded is not None:
    if not uploaded.name.lower().endswith(".pdf"):
        st.error("Please upload a PDF file.")
    else:
        # Make client
        try:
            client = make_client(AZURE_DI_ENDPOINT or "", AZURE_DI_KEY or "")
        except Exception as e:
            st.error(f"Failed to create Azure DI client: {e}")
            st.stop()

        pdf_bytes = uploaded.read()
        if not pdf_bytes:
            st.error("Uploaded file is empty. Please re-upload the PDF.")
            st.stop()

        # Hash to prevent double-charging on re-runs
        fid = file_hash(pdf_bytes)

        with st.spinner("Analyzing with Azure Document Intelligence (prebuilt-read)..."):
            try:
                result = analyze_pdf_bytes(client, pdf_bytes)
            except Exception as e:
                st.error(f"Azure DI analyze failed: {e}")
                st.stop()

        # Pages (prefer DI, fallback to 1)
        pages = len(getattr(result, "pages", []) or [])
        if pages <= 0:
            pages = 1
        st.success(f"Extracted text from **{pages} page(s)**.")

        # Charge (only once per file hash)
        try:
            charged = charge_credits_once(fid, pages, uploaded.name)
            if charged > 0:
                st.toast(f"Charged {charged} credits for {pages} page(s).", icon="✅")
        except RuntimeError as e:
            st.error(str(e))
            st.stop()

        # Build DOCX
        with st.spinner("Building DOCX..."):
            try:
                docx_bytes = result_to_docx_bytes(
                    result,
                    insert_page_breaks=add_page_breaks,
                    show_conf=include_confidence
                )
            except Exception as e:
                st.error(f"Failed to create DOCX: {e}")
                st.stop()

        st.download_button(
            label="⬇️ Download .docx",
            data=docx_bytes,
            file_name=(uploaded.name.rsplit(".", 1)[0] + ".docx"),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        with st.expander("Preview extracted text (first ~2,000 chars)"):
            preview = getattr(result, "content", "")
            st.text(preview[:2000] + ("..." if len(preview) > 2000 else ""))

else:
    st.info("Upload a PDF to begin.")

# =========================
# FOOTER
# =========================
st.caption(
    "Credits are session-scoped in this demo and locked for users. Only admins can refill using the PIN from st.secrets. "
    f"Pricing: {PRICE_PER_PAGE_CREDITS} credits (₹{PRICE_PER_PAGE_CREDITS}) per page."
)
